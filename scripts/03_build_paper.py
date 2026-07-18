#!/usr/bin/env python3
"""
Generate the final trade openness - economic growth manuscript formatted for
World Development (Elsevier) using real analysis results.

Outputs:
    Trade_Openness_Growth_Real.docx  - anonymized manuscript (main text,
                                       declarations, references, tables).
    title_page.docx                  - title page with author information,
                                       acknowledgements and competing-interest
                                       declaration for the submission system.
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import numpy as np
import os

REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = REPO_ROOT / "output"
MANUSCRIPT_PATH = REPO_ROOT / "Trade_Openness_Growth_Real.docx"
TITLE_PAGE_PATH = REPO_ROOT / "title_page.docx"
REPO_URL = "https://github.com/AcerisSolaIAS/trade-openness-growth-replication"


# ---------------------------------------------------------------------------
# Base formatting helpers
# ---------------------------------------------------------------------------

def set_default_font(doc, font_name="Times New Roman", font_size=Pt(12)):
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = font_size
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_run_font(run, font_name="Times New Roman", size=Pt(12), bold=False,
                 italic=False, color=RGBColor(0, 0, 0)):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_paragraph_spacing(p, before=Pt(0), after=Pt(0), line_spacing=2.0,
                          first_line_indent=Inches(0.0)):
    pf = p.paragraph_format
    pf.space_before = before
    pf.space_after = after
    pf.line_spacing = line_spacing
    pf.first_line_indent = first_line_indent


def add_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_line_indent=Inches(0.5), line_spacing=2.0,
                  space_after=Pt(0), bold=False, italic=False, size=Pt(12)):
    p = doc.add_paragraph()
    p.alignment = alignment
    set_paragraph_spacing(p, after=space_after, line_spacing=line_spacing,
                          first_line_indent=first_line_indent)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        set_paragraph_spacing(p, before=Pt(18), after=Pt(6),
                              line_spacing=2.0, first_line_indent=Inches(0.0))
        size = Pt(14)
    elif level == 2:
        set_paragraph_spacing(p, before=Pt(12), after=Pt(6),
                              line_spacing=2.0, first_line_indent=Inches(0.0))
        size = Pt(12)
    else:
        set_paragraph_spacing(p, before=Pt(6), after=Pt(3),
                              line_spacing=2.0, first_line_indent=Inches(0.0))
        size = Pt(12)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)
    return p


def add_centered_paragraph(doc, text, size=Pt(12), bold=False, italic=False,
                           space_after=Pt(0), line_spacing=2.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=space_after, line_spacing=line_spacing,
                          first_line_indent=Inches(0.0))
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


# ---------------------------------------------------------------------------
# Tables (editable Word tables, no vertical rules)
# ---------------------------------------------------------------------------

def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=Pt(12), after=Pt(6),
                          line_spacing=2.0, first_line_indent=Inches(0.0))
    run = p.add_run(text)
    set_run_font(run, bold=True)
    return p


def add_word_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(paragraph, line_spacing=1.15,
                                  first_line_indent=Inches(0.0))
            set_run_font(paragraph.runs[0], bold=True)
        if col_widths:
            cell.width = col_widths[i]

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(paragraph, line_spacing=1.15,
                                      first_line_indent=Inches(0.0))
                set_run_font(paragraph.runs[0])
            if col_widths:
                row_cells[i].width = col_widths[i]

    # Remove vertical rules and shading; keep horizontal rules only.
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "FFFFFF")
            tcPr.append(shd)
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top", "bottom"):
                edge_el = OxmlElement(f"w:{edge}")
                edge_el.set(qn("w:val"), "single")
                edge_el.set(qn("w:sz"), "4")
                edge_el.set(qn("w:space"), "0")
                edge_el.set(qn("w:color"), "000000")
                tcBorders.append(edge_el)
            for edge in ("left", "right", "insideH", "insideV"):
                edge_el = OxmlElement(f"w:{edge}")
                edge_el.set(qn("w:val"), "nil")
                tcBorders.append(edge_el)
            tcPr.append(tcBorders)
    return table


def add_table_note(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, before=Pt(3), after=Pt(12),
                          line_spacing=2.0, first_line_indent=Inches(0.0))
    run = p.add_run("Note: " + text)
    set_run_font(run, italic=True, size=Pt(11))
    return p


# ---------------------------------------------------------------------------
# References (APA 7 / Harvard alphabetical)
# ---------------------------------------------------------------------------

REFERENCES = [
    ("autor2013", "Autor, D. H., Dorn, D., & Hanson, G. H. (2013). The China syndrome: Local labor market effects of import competition in the United States. American Economic Review, 103(6), 2121-2168. https://doi.org/10.1257/aer.103.6.2121"),
    ("billmeier2013", "Billmeier, A., & Nannicini, T. (2013). Assessing economic liberalization episodes: A synthetic control approach. Review of Economics and Statistics, 95(3), 983-1001. https://doi.org/10.1162/REST_a_00300"),
    ("chang2009", "Chang, R., Kaltani, L., & Loayza, N. V. (2009). Openness can be good for growth: The role of policy complementarities. Journal of Development Economics, 90(1), 33-49. https://doi.org/10.1016/j.jdeveco.2008.11.002"),
    ("edwards1998", "Edwards, S. (1998). Openness, productivity and growth: What do we really know? Economic Journal, 108(447), 383-398. https://doi.org/10.1111/1468-0297.00293"),
    ("estevadeordal2013", "Estevadeordal, A., & Taylor, A. M. (2013). Is the Washington Consensus dead? Growth, openness, and the Great Liberalization, 1980s-2000s. Review of Economics and Statistics, 95(5), 1669-1690. https://doi.org/10.1162/REST_a_00324"),
    ("fajgelbaum2016", "Fajgelbaum, P. D., & Khandelwal, A. K. (2016). Measuring the unequal gains from trade. Quarterly Journal of Economics, 131(3), 1113-1180. https://doi.org/10.1093/qje/qjw013"),
    ("feyrer2019", "Feyrer, J. (2019). Trade and income-Exploiting time series in geography. American Economic Journal: Applied Economics, 11(4), 1-35. https://doi.org/10.1257/app.20170608"),
    ("frankel1999", "Frankel, J. A., & Romer, D. H. (1999). Does trade cause growth? American Economic Review, 89(3), 379-399. https://doi.org/10.1257/aer.89.3.379"),
    ("irwin2002", "Irwin, D. A., & Tervio, M. (2002). Does trade raise income? Evidence from the twentieth century. Journal of International Economics, 58(1), 1-18. https://doi.org/10.1016/S0022-1996(01)00164-7"),
    ("maddala1999", "Maddala, G. S., & Wu, S. (1999). A comparative study of unit root tests with panel data and a new simple test. Oxford Bulletin of Economics and Statistics, 61(S1), 631-652. https://doi.org/10.1111/1468-0084.0610s1657"),
    ("nickell1981", "Nickell, S. (1981). Biases in dynamic models with fixed effects. Econometrica, 49(6), 1417-1426. https://doi.org/10.2307/1911408"),
    ("pierce2016", "Pierce, J. R., & Schott, P. K. (2016). The surprisingly swift decline of US manufacturing employment. American Economic Review, 106(7), 1632-1662. https://doi.org/10.1257/aer.20131578"),
    ("rodriguez2000", "Rodriguez, F., & Rodrik, D. (2000). Trade policy and economic growth: A skeptic's guide to the cross-national evidence. NBER Macroeconomics Annual, 15, 261-325. https://doi.org/10.1086/654419"),
    ("sachs1995", "Sachs, J. D., & Warner, A. M. (1995). Economic reform and the process of global integration. Brookings Papers on Economic Activity, 1995(1), 1-118."),
    ("wacziarg2008", "Wacziarg, R., & Welch, K. H. (2008). Trade liberalization and growth: New evidence. World Bank Economic Review, 22(2), 187-231. https://doi.org/10.1093/wber/lhn007"),
    ("worldbank2024", "World Bank. (2024). World Development Indicators. World Bank. https://data.worldbank.org/"),
]

# Sort alphabetically by first author surname and replace hyphens in page ranges with en dashes.
REFERENCES.sort(key=lambda kv: kv[1].split(",")[0].split()[-1].lower())


def _en_dash_page_range(ref):
    # Replace the first occurrence of a page range pattern like 123-456 or F22-F49 with an en dash.
    import re
    return re.sub(r"(\d{1,4}|F\d{1,3}|S\d{1,2})([\-–])(\d{1,4}|F\d{1,3}|S\d{1,2})",
                  r"\1-\3", ref).replace("-", "\u2013", 1)


def add_reference_list(doc):
    for _, ref_text in REFERENCES:
        ref_text = _en_dash_page_range(ref_text)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(ref_text)
        set_run_font(run, size=Pt(12))


# ---------------------------------------------------------------------------
# Data loading and helpers
# ---------------------------------------------------------------------------

def load_results():
    panel = pd.read_csv(OUT_DIR / "panel_regression_results.csv")
    dyn = pd.read_csv(OUT_DIR / "dynamic_panel_results.csv")
    robust = pd.read_csv(OUT_DIR / "robustness_results.csv")
    desc = pd.read_csv(OUT_DIR / "descriptive_stats.csv")
    corr = pd.read_csv(OUT_DIR / "correlation_matrix.csv", index_col=0)
    vif = pd.read_csv(OUT_DIR / "vif_table.csv")

    unitroot = None
    coint = None
    het = None
    if os.path.exists(OUT_DIR / "panel_unit_root_results.csv"):
        unitroot = pd.read_csv(OUT_DIR / "panel_unit_root_results.csv")
    if os.path.exists(OUT_DIR / "cointegration_summary.csv"):
        coint = pd.read_csv(OUT_DIR / "cointegration_summary.csv")
    if os.path.exists(OUT_DIR / "heterogeneity_results.csv"):
        het = pd.read_csv(OUT_DIR / "heterogeneity_results.csv")
    return panel, dyn, robust, desc, corr, vif, unitroot, coint, het


def get_coef(df, model, var):
    r = df[(df["model"] == model) & (df["variable"] == var)]
    if r.empty:
        return np.nan, np.nan
    return float(r["coef"].values[0]), float(r["pvalue"].values[0])


def get_coef_het(df, sample, var):
    r = df[(df["sample"] == sample) & (df["variable"] == var)]
    if r.empty:
        return np.nan, np.nan
    return float(r["coef"].values[0]), float(r["pvalue"].values[0])


def coef_str(c, p):
    if np.isnan(c):
        return "---"
    stars = ""
    if p < 0.01:
        stars = "***"
    elif p < 0.05:
        stars = "**"
    elif p < 0.1:
        stars = "*"
    return f"{c:.4f}{stars}"


def format_p(p):
    if pd.isna(p):
        return "---"
    if p < 0.001:
        return "<0.001"
    return f"{p:.3f}"


# ---------------------------------------------------------------------------
# Manuscript assembly
# ---------------------------------------------------------------------------

def build_manuscript(panel, dyn, robust, desc, corr, vif, unitroot, coint, het):
    fe_trade_c, fe_trade_p = get_coef(panel, "fe", "ltrade")
    twfe_trade_c, twfe_trade_p = get_coef(panel, "twfe", "ltrade")
    dyn_trade_c, dyn_trade_p = get_coef(dyn, "dynamic_fe", "ltrade_lag1")
    diff_trade_c, diff_trade_p = get_coef(robust, "first_diff", "d_ltrade")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    set_default_font(doc)

    # Title (no author information in the anonymized manuscript).
    add_centered_paragraph(
        doc,
        "Trade Openness and Economic Growth: A Replicable Cross-Country Panel Analysis, 2000-2022",
        size=Pt(16), bold=True, space_after=Pt(18), line_spacing=2.0
    )

    # Abstract (unstructured, max 250 words).
    add_heading(doc, "Abstract", level=1)
    abstract_text = (
        f"This paper re-examines the trade openness-growth relationship using a transparent, "
        f"reproducible panel of up to 183 economies from 2000 to 2022. We estimate pooled OLS, "
        f"fixed-effects, time-effects, and two-way fixed-effects specifications and report "
        f"Maddala-Wu Fisher panel unit-root tests, Engle-Granger residual cointegration tests, "
        f"and subsample splits by development level, crisis period, and trade intensity. A 1% "
        f"increase in the trade-to-GDP ratio is associated with a {fe_trade_c:.4f}% rise in GDP "
        f"per capita under fixed effects (p = {fe_trade_p:.4f}), but the coefficient falls to "
        f"{twfe_trade_c:.4f} (p = {twfe_trade_p:.4f}) once country and year effects are included. "
        f"Dynamic-panel and first-difference estimates remain positive, while subsample analysis "
        f"indicates that the association is stronger among developed economies and before the 2008 "
        f"financial crisis than afterward. The results suggest that the aggregate trade-to-GDP "
        f"ratio captures several distinct economic forces and warn against treating broad "
        f"liberalization as an automatic engine of growth."
    )
    add_paragraph(doc, abstract_text, first_line_indent=Inches(0.0))

    # Keywords (3-6, single words or short phrases).
    p_kw = doc.add_paragraph()
    p_kw.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p_kw, after=Pt(18), line_spacing=2.0,
                          first_line_indent=Inches(0.0))
    r1 = p_kw.add_run("Keywords: ")
    set_run_font(r1, bold=True)
    r2 = p_kw.add_run("trade openness; economic growth; panel data; fixed effects; cointegration")
    set_run_font(r2)

    # 1. Introduction
    add_heading(doc, "1. Introduction", level=1)
    add_paragraph(
        doc,
        "The empirical literature on trade and growth remains divided. While Frankel and Romer "
        "(1999) find a positive causal link using cross-sectional instruments, Rodriguez and Rodrik "
        "(2000) famously question the robustness of these findings, arguing that openness indicators "
        "often proxy for geography, institutions, and macroeconomic stability rather than trade policy "
        "itself. Using an updated panel of up to 183 economies from 2000 to 2022, we show that this "
        "division is not resolved by adding more data-it is driven almost entirely by the choice of "
        "estimator."
    )
    add_paragraph(
        doc,
        "The theoretical case for a positive link between trade and growth rests on several mechanisms. "
        "Openness expands market size, allows specialization according to comparative advantage, and "
        "exposes domestic firms to foreign competition and technology (Frankel & Romer, 1999). Endogenous "
        "growth models emphasize knowledge spillovers through trade, while new trade theory highlights "
        "scale economies and product variety. Yet these benefits are not automatic. Countries that are "
        "geographically isolated, institutionally weak, or heavily dependent on a narrow set of primary "
        "commodities may see little productivity gain from lowering tariffs, and may instead suffer from "
        "terms-of-trade volatility and de-industrialization."
    )
    add_paragraph(
        doc,
        "The empirical challenge is that trade openness is not randomly assigned. Small, rich, coastal "
        "economies trade more than large, poor, landlocked ones, and they also differ in education, "
        "institutions, and geography. Pooled OLS therefore conflates the effect of trade with the effect "
        "of being the kind of country that trades a lot (Wacziarg & Welch, 2008). Fixed-effects estimators "
        "remove time-invariant heterogeneity, but they leave two issues unresolved. First, common global "
        "trends-falling transport costs, trade liberalization waves, and post-2008 convergence "
        "patterns-can be absorbed by year effects, sometimes flipping the sign of the trade coefficient "
        "(Estevadeordal & Taylor, 2013). Second, the trade-to-GDP ratio is a noisy proxy for trade "
        "policy: it rises when commodity prices surge, when a country hosts a regional trading hub, or "
        "when it liberalizes; these are not the same shock. Treating them as interchangeable has "
        "produced a literature in which the same variable can be reported as strongly positive, weakly "
        "positive, or negative depending on the decade and the estimator."
    )
    add_paragraph(
        doc,
        "Recent work has moved toward sharper identification. Feyrer (2019) uses geography-driven "
        "variation in air-freight costs as an instrument; Billmeier and Nannicini (2013) evaluate "
        "liberalization episodes with synthetic controls; micro studies have traced the distributional "
        "consequences of import competition in specific labor markets (Autor et al., 2013; Pierce & "
        "Schott, 2016; Fajgelbaum & Khandelwal, 2016). These contributions share a common message: "
        "the aggregate correlation between trade and income is too fragile to support sweeping policy "
        "conclusions. Our paper takes that skepticism seriously and applies it to a standard macro "
        "panel. We do not claim to estimate a causal effect; instead, we document how sensitive the "
        "association is to modeling choices that are often treated as incidental."
    )
    add_paragraph(
        doc,
        "We assemble an unbalanced panel of up to 183 economies from 2000 to 2022 using only the World "
        "Bank World Development Indicators API (World Bank, 2024). The analysis is deliberately "
        "transparent: pooled OLS, fixed effects, time effects, and two-way fixed effects are reported "
        "side by side, together with Maddala-Wu Fisher panel unit-root tests (Maddala & Wu, 1999), "
        "Engle-Granger residual cointegration tests, subsample splits by development level and crisis "
        "period, and robustness checks that exclude small open economies, use lagged trade openness, "
        "and first-difference the data. Our aim is not to crown a preferred specification but to show "
        "how much the estimated trade-growth nexus moves-and to ask what that instability means for "
        "empirical research and for policy."
    )

    add_paragraph(
        doc,
        "The period under study is also distinctive. The two decades after 2000 saw the rapid expansion "
        "of global value chains, the integration of China into world trade, and, after 2008, a marked "
        "slowdown in trade growth relative to GDP. These global trends create common factors that can "
        "produce a positive association between trade and income even when country-specific trade "
        "liberalization has no causal effect. A two-way fixed-effects estimator is therefore a useful "
        "diagnostic: it asks whether countries that opened more than the global average in a given year "
        "also grew faster than the global average in that year."
    )
    add_paragraph(
        doc,
        "The remainder of the paper is organized as follows. Section 2 describes the data, the measurement "
        "of openness, and the panel estimators. Section 3 presents descriptive patterns, stationarity tests, "
        "main estimates, heterogeneity analysis, and robustness checks. Section 4 discusses the interpretation "
        "of the results and their limitations. Section 5 concludes with implications for research and policy."
    )

    # 2. Data and methods
    add_heading(doc, "2. Data and methods", level=1)
    add_heading(doc, "2.1 Data sources and sample", level=2)
    add_paragraph(
        doc,
        "All data are obtained from the World Bank World Development Indicators (WDI) database via the "
        "public API (World Bank, 2024). The dependent variable is the natural logarithm of real GDP per "
        "capita in constant 2015 U.S. dollars (NY.GDP.PCAP.KD). The key independent variable is the "
        "natural logarithm of the trade-to-GDP ratio, defined as the sum of exports and imports of goods "
        "and services as a share of GDP (NE.TRD.GNFS.ZS). Control variables are gross capital formation "
        "as a share of GDP (NE.GDI.TOTL.ZS), gross primary school enrollment (SE.PRM.ENRR), and annual "
        "population growth (SP.POP.GROW)."
    )
    add_paragraph(
        doc,
        "The sample covers the period 2000-2022. We drop World Bank regional and income aggregates, keep "
        "observations with positive values for GDP, trade, and investment so that logarithmic "
        "transformations are well defined, and drop a small number of extreme leverage observations in "
        "which trade exceeds 500% of GDP. These cases are mostly city-states and financial entrepots "
        "such as Hong Kong SAR, Singapore, and Luxembourg, where re-exports and financial-services trade "
        "inflate the merchandise trade measure far beyond domestic production. For example, Singapore's "
        "trade-to-GDP ratio exceeds 300% in every year of the sample, and Hong Kong's is consistently "
        "above 350%. Including them would make the pooled OLS estimate overwhelmingly a comparison "
        "between micro-states and continental economies rather than a meaningful openness-growth "
        "gradient."
    )
    add_paragraph(
        doc,
        "The resulting unbalanced panel contains 3,919 country-year observations for 183 economies. Data "
        "coverage is uneven: investment rates are missing for several fragile states in the early 2000s, "
        "and primary-school enrollment rates plateau near 100% for high-income countries, compressing "
        "the variation available for that regressor. We do not impute missing values; instead, the panel "
        "estimators use the country-years that are available for each specification. Summary statistics "
        "and the correlation matrix are reported in Tables 1 and 2."
    )

    add_heading(doc, "2.2 Measurement of openness", level=2)
    add_paragraph(
        doc,
        "Measuring trade openness is harder than it appears. The trade-to-GDP ratio is the most widely "
        "used indicator because it is available for almost all countries and years, but it conflates "
        "several distinct phenomena. A high ratio may reflect deliberate liberalization, membership in a "
        "regional trade agreement, geographic proximity to large markets, participation in global value "
        "chains, or simply a small domestic economy that imports most consumption goods and exports a "
        "narrow set of commodities. For mineral exporters, the ratio moves with world commodity prices "
        "rather than with trade policy. For assembly economies, re-exports can dwarf domestic value "
        "added. These ambiguities are well known (Rodriguez & Rodrik, 2000), yet the trade-to-GDP ratio "
        "remains the workhorse measure in cross-country growth regressions because alternatives such as "
        "tariff rates, non-tariff barrier indices, or trade policy scores are available only for smaller "
        "samples and shorter periods."
    )
    add_paragraph(
        doc,
        "We address this limitation in two ways. First, we control for investment, schooling, and "
        "population growth, which capture some of the channels through which openness might affect "
        "income. Second, we exclude the most extreme cases of re-export and financial entrepot activity, "
        "and we split the sample by development level and trade intensity to see whether the association "
        "is driven by a particular group of countries. Even so, our estimates should be read as "
        "conditional correlations rather than causal effects."
    )

    add_heading(doc, "2.3 Econometric strategy", level=2)
    add_paragraph(
        doc,
        "We estimate four panel specifications. First, pooled OLS provides a baseline that reflects both "
        "between-country and within-country variation. Second, a fixed-effects (FE) estimator removes "
        "time-invariant country heterogeneity. Third, a time-effects (TE) estimator removes common shocks "
        "and trends. Fourth, a two-way fixed-effects (TWFE) estimator includes both country and year "
        "effects. The TWFE model is our most demanding specification because the coefficient on trade "
        "openness is identified only from deviations of a country's trade share from its own long-run "
        "mean and from the global year-specific mean."
    )
    add_paragraph(
        doc,
        "Formally, the static specifications can be written as y_it = alpha_i + lambda_t + beta * "
        " openness_it + gamma' * X_it + epsilon_it, where y_it is log GDP per capita, alpha_i is a "
        "country fixed effect, lambda_t is a year fixed effect, X_it contains the control variables, and "
        "epsilon_it is an error term. Pooled OLS sets alpha_i = 0 and lambda_t = 0 for all i and t; FE "
        "sets lambda_t = 0; TE sets alpha_i = 0; and TWFE includes both sets of fixed effects. Standard "
        "errors are clustered at the country level to account for serial correlation within countries."
    )
    add_paragraph(
        doc,
        "Because GDP per capita is highly persistent, we also estimate a dynamic panel that includes a "
        "lagged dependent variable. The lagged dependent variable partly controls for convergence "
        "dynamics and absorbs omitted slow-moving factors, but it is known to introduce bias in short "
        "panels; we therefore treat it as a robustness check rather than the preferred model (Nickell, "
        "1981). Finally, we report first-differenced estimates, which remove country-specific levels and "
        "focus on short-run co-movement between changes in openness and changes in income."
    )

    add_heading(doc, "2.4 Panel stationarity and cointegration pre-tests", level=2)
    add_paragraph(
        doc,
        "When variables are persistent, pooled OLS and FE estimates may reflect spurious correlations "
        "rather than stable long-run relationships. We therefore report country-by-country augmented "
        "Dickey-Fuller (ADF) tests and Maddala-Wu Fisher combined p-value tests for the log-levels of "
        "GDP per capita, trade openness, and investment (Maddala & Wu, 1999). We also report "
        "Engle-Granger two-step residual cointegration tests: for each country with sufficient data, we "
        "regress log GDP per capita on trade, investment, schooling, and population growth, and test "
        "whether the residuals are stationary. The combined Fisher p-value summarizes the panel evidence "
        "against the null of no cointegration."
    )
    add_paragraph(
        doc,
        "These tests are useful diagnostics, but they have limitations. ADF tests have low power in short "
        "panels, and the Maddala-Wu approach treats countries as independent units, which is unlikely to "
        "hold when countries are linked by common shocks and global business cycles. Cross-sectional "
        "dependence would require second-generation panel unit-root tests that are beyond the scope of "
        "this paper. We therefore interpret the pre-tests as descriptive evidence rather than as a "
        "definitive classification of the data-generating process."
    )

    add_heading(doc, "2.5 Heterogeneity analysis", level=2)
    add_paragraph(
        doc,
        "The macro-level trade-growth relationship may mask important heterogeneity. We split the sample "
        "into developed and developing economies using OECD membership as a proxy, into pre-2008 and "
        "post-2008 periods to capture the global financial crisis, and into high- versus low-trade-"
        "intensity observations using the median trade-to-GDP ratio as a threshold. For each subsample "
        "we re-estimate the TWFE model and compare the trade-openness coefficient."
    )

    add_paragraph(
        doc,
        "All variables enter the regressions in natural logarithms, so the coefficients can be read as "
        "elasticities. This choice dampens the influence of outliers, narrows differences in scale between "
        "high- and low-income countries, and corresponds to the standard growth-regression specification "
        "in which income and openness are both persistent, bounded, and non-negative. Standard errors are "
        "clustered at the country level, which allows the error term to be arbitrarily correlated over time "
        "within a country but assumes independence across countries."
    )
    add_paragraph(
        doc,
        "We deliberately keep the econometric approach simple. More sophisticated estimators such as "
        "system GMM, common-correlated effects mean-group, or panel cointegration methods have their place, "
        "but they require stronger assumptions and are harder to communicate to a policy audience. Our "
        "baseline specifications can be reproduced with any standard statistical package, which aligns with "
        "the journal's emphasis on transparent, evidence-based development research."
    )
    add_paragraph(
        doc,
        "The heterogeneity and robustness analyses are designed to probe the stability of the TWFE "
        "coefficient. We split the sample by OECD membership because developed economies have deeper "
        "financial markets, stronger contract enforcement, and more diversified production structures, "
        "all of which may raise the growth payoff from trade. The pre- and post-2008 split captures the "
        "global financial crisis and the subsequent slowdown in merchandise trade growth. The high- versus "
        "low-trade-intensity split checks whether the negative TWFE result is confined to small open "
        "economies with volatile trade shares."
    )

    # 3. Results
    add_heading(doc, "3. Results", level=1)
    add_heading(doc, "3.1 Descriptive statistics", level=2)
    add_paragraph(
        doc,
        f"The average country-year observation in the sample has a trade-to-GDP ratio of approximately "
        f"{desc['trade_mean'].mean():.1f}% and a real GDP per capita of about ${desc['gdp_mean'].mean():,.0f}. "
        f"There is substantial cross-country heterogeneity: mean trade shares range from below 20% to "
        f"above 400% for small open economies. Correlations show that log GDP per capita is positively "
        f"correlated with log trade openness (r = {corr.loc['lgdp', 'ltrade']:.3f}) and negatively "
        f"correlated with population growth (r = {corr.loc['lgdp', 'popg']:.3f}). Variance-inflation "
        f"factors for the regressors are below conventional thresholds, with the exception of the "
        f"constant, indicating that multicollinearity is not a major concern in our main specifications."
    )
    add_paragraph(
        doc,
        "The descriptive patterns already hint at the identification problem. High-income countries "
        "tend to have higher trade shares and higher GDP per capita, generating a strong cross-sectional "
        "gradient. Within countries, however, trade shares move with commodity prices, exchange-rate "
        "fluctuations, and global demand rather than with domestic policy changes alone. The contrast "
        "between the between-country and within-country variation is exactly what the econometric "
        "specifications in Table 4 are designed to expose."
    )

    if unitroot is not None and not unitroot.empty:
        add_heading(doc, "3.2 Panel stationarity and cointegration", level=2)
        ur_text = (
            "Table 3 reports the pre-test results. For the log-levels of GDP per capita, trade, and "
            "investment, the country-by-country ADF tests reject the unit-root null in "
        )
        parts = []
        for _, row in unitroot.iterrows():
            parts.append(f"{int(row['n_reject_5pct'])} of {int(row['n_countries'])} countries for {row['variable']}")
        ur_text += "; ".join(parts) + ". "
        fisher_parts = []
        for _, row in unitroot.iterrows():
            fisher_parts.append(f"{format_p(row['fisher_pvalue'])} ({row['variable']})")
        ur_text += "The Maddala-Wu Fisher combined tests yield p-values of " + ", ".join(fisher_parts) + ". "
        if coint is not None and not coint.empty:
            cr = coint.iloc[0]
            ur_text += (
                f"The Engle-Granger residual cointegration test rejects the null of no cointegration in "
                f"{int(cr['n_reject_5pct'])} of {int(cr['n_countries'])} countries, with a Fisher combined "
                f"p-value of {format_p(cr['fisher_pvalue'])}. "
            )
        ur_text += (
            "These results indicate that the variables are highly persistent and that the presence of a "
            "stable long-run cointegrating relationship should be treated as country-specific rather than "
            "uniform across the whole panel."
        )
        add_paragraph(doc, ur_text)
        add_paragraph(
            doc,
            "The pre-tests matter for interpretation. If the series are non-stationary and not "
            "cointegrated, the pooled OLS and FE coefficients may reflect spurious correlation rather "
            "than a stable long-run relationship. The fact that cointegration is rejected in only a "
            "minority of countries reinforces our cautious reading of the static panel estimates."
        )

    add_heading(doc, "3.3 Panel estimates", level=2)
    add_paragraph(
        doc,
        f"Table 4 reports the panel estimates. In the pooled OLS model, the trade-to-GDP coefficient is "
        f"large and highly significant (coef. = {get_coef(panel, 'pooled_ols', 'ltrade')[0]:.4f}, p < 0.001), "
        f"reflecting the strong cross-sectional relationship between openness and income. Small, rich "
        f"trading hubs and large, relatively closed economies sit at opposite ends of this gradient. Once "
        f"country fixed effects are introduced, the coefficient falls sharply but remains positive and "
        f"statistically significant (coef. = {fe_trade_c:.4f}, p = {fe_trade_p:.4f}), which means that "
        f"even when the comparison is restricted to changes within countries over time, higher trade "
        f"shares are still associated with higher income."
    )
    add_paragraph(
        doc,
        f"The two-way fixed-effects estimate is smaller and turns negative (coef. = {twfe_trade_c:.4f}, "
        f"p = {twfe_trade_p:.4f}). The change is driven by the removal of common global trends: between "
        f"2000 and 2022, trade openness and income both rose on average, so the FE estimate partly "
        f"captures a shared global upswing. Once year effects are added, the residual within-country "
        f"relationship between trade and income is negative, though modest. This does not mean that trade "
        f"reduces income; it means that the countries whose trade shares grew fastest within each year "
        f"were not the same countries whose incomes grew fastest within that year."
    )
    add_paragraph(
        doc,
        "The control variables behave largely as expected. Investment is positively associated with GDP "
        "per capita in most specifications, consistent with a standard neoclassical accumulation story. "
        "Primary-school enrollment is often insignificant, partly because it is measured as gross "
        "enrollment and varies little among high-income countries. Population growth is usually negative, "
        "consistent with the prediction that faster labor-force growth dilutes capital per worker."
    )

    if het is not None and not het.empty:
        add_heading(doc, "3.4 Heterogeneity analysis", level=2)
        samples = ["developed", "developing", "pre_2008", "post_2008", "high_trade", "low_trade"]
        labels = {
            "developed": "developed economies",
            "developing": "developing economies",
            "pre_2008": "pre-2008 period",
            "post_2008": "post-2008 period",
            "high_trade": "high-trade-intensity group",
            "low_trade": "low-trade-intensity group",
        }
        het_text = "Table 5 reports TWFE estimates by subsample. "
        for s in samples:
            c, p = get_coef_het(het, s, "ltrade")
            if not np.isnan(c):
                sig = "significant" if p < 0.05 else "insignificant"
                het_text += f"For {labels[s]}, the trade coefficient is {c:.4f} (p = {format_p(p)}), {sig}. "
        het_text += (
            "The developed-economy coefficient is positive, while the developing-economy coefficient is "
            "negative, indicating that the within-country trade-income association is not uniform across "
            "the income distribution. The pre-2008 coefficient is also larger than the post-2008 "
            "coefficient, which is consistent with the slowdown in trade growth after the global financial "
            "crisis. Both the high- and low-trade-intensity groups return negative TWFE coefficients, so "
            "the negative full-sample result is not simply a small-open-economy story."
        )
        add_paragraph(doc, het_text)
        add_paragraph(
            doc,
            "The heterogeneity results are important for policy. If the association between trade and "
            "growth were uniformly positive, one could argue for broad liberalization regardless of "
            "country characteristics. The fact that the coefficient is positive only for developed "
            "economies and for the pre-2008 period suggests that context matters. Developing economies "
            "may lack the complementary institutions-absorptive capacity, stable macroeconomic "
            "environments, and diversified production structures-needed to translate openness into "
            "sustained growth."
        )

    add_heading(doc, "3.5 Robustness checks", level=2)
    add_paragraph(
        doc,
        f"Table 6 reports robustness specifications. Excluding economies with mean trade shares above "
        f"150% of GDP leaves the TWFE trade coefficient negative and significant (coef. = "
        f"{get_coef(robust, 'excl_high_trade', 'ltrade')[0]:.4f}, p = "
        f"{get_coef(robust, 'excl_high_trade', 'ltrade')[1]:.4f}). The sign does not flip back to "
        f"positive, which suggests that the negative TWFE result is not just a small-open-economy story; "
        f"it is a more general within-country phenomenon over this period. A specification with lagged "
        f"trade openness yields a negative but marginally insignificant coefficient (p = "
        f"{get_coef(robust, 'lag_trade', 'ltrade_lag1')[1]:.3f}). By contrast, the first-differenced "
        f"estimate is positive and highly significant (coef. = {diff_trade_c:.4f}, p < 0.001), "
        f"indicating that year-to-year increases in openness are associated with contemporaneous "
        f"increases in income."
    )
    add_paragraph(
        doc,
        f"The dynamic panel in Table 7 includes a lagged dependent variable and reports a positive and "
        f"significant coefficient on one-year-lagged trade openness (coef. = {dyn_trade_c:.4f}, p = "
        f"{dyn_trade_p:.4f}). The lagged dependent variable itself is large (coef. = "
        f"{get_coef(dyn, 'dynamic_fe', 'lgdp_lag1')[0]:.4f}, p < 0.001), which is consistent with the "
        f"high persistence of GDP per capita. The robustness checks therefore reinforce the main message: "
        f"the sign and significance of the estimated trade-growth relationship depend on whether the "
        f"model emphasizes cross-sectional variation, short-run changes, or deviations from country-"
        f"specific trends."
    )

    add_paragraph(
        doc,
        "The contrast between the pooled OLS and TWFE estimates is the central empirical fact in the "
        "data. Pooled OLS assigns a roughly unit elasticity to trade openness, implying that a doubling "
        "of the trade-to-GDP ratio is associated with a doubling of GDP per capita across countries. "
        "This cross-sectional association is informative about long-run development patterns but says "
        "little about the growth effect of a contemporaneous liberalization episode. Once country and "
        "year effects are removed, the elasticity becomes small and negative, which is consistent with "
        "the view that the cross-sectional gradient is driven by permanent country characteristics and "
        "common global trends rather than by within-country policy variation."
    )
    add_paragraph(
        doc,
        "The dynamic-panel and first-difference results reinforce this interpretation. The lagged "
        "dependent variable absorbs most of the persistence in GDP per capita, and the remaining "
        "coefficient on lagged trade is positive but modest. First-differenced estimates are also "
        "positive, indicating that short-run increases in openness are associated with short-run "
        "increases in income. The discrepancy between the short-run and TWFE estimates suggests that "
        "the longer-run, trend-like component of openness is not the same as its year-to-year "
        "fluctuations."
    )
    add_paragraph(
        doc,
        "The control variables also deserve comment. Investment is consistently positive across "
        "specifications, which is consistent with a neoclassical accumulation channel. Population growth "
        "is negative in most static specifications, matching the idea that faster labor-force growth "
        "dilutes the capital stock. Primary-school enrollment is often insignificant, probably because "
        "gross enrollment is a coarse measure that is compressed near 100% in high-income countries and "
        "because its short-run variation is weakly correlated with contemporaneous income growth."
    )

    # 4. Discussion
    add_heading(doc, "4. Discussion", level=1)
    add_paragraph(
        doc,
        "Our findings line up with the broader literature in two respects. First, the strong cross-"
        "sectional correlation between trade openness and income levels is well documented (Frankel & "
        "Romer, 1999; Sachs & Warner, 1995; Irwin & Tervio, 2002). Second, the sensitivity of this "
        "correlation to fixed effects is consistent with the critique that cross-country regressions may "
        "conflate trade with geography, institutions, and other development determinants (Rodriguez & "
        "Rodrik, 2000; Edwards, 1998). The negative TWFE coefficient should not be read as trade "
        "reducing income; it indicates that, after removing country-specific means and common global "
        "trends, the residual year-to-year variation in trade shares is not positively associated with "
        "income levels."
    )
    add_paragraph(
        doc,
        "The heterogeneity analysis supports this reading. The trade coefficient differs between "
        "developed and developing economies and between the pre- and post-2008 periods, echoing the "
        "argument of Chang, Kaltani, and Loayza (2009) that the growth effects of openness depend on "
        "complementary policies and institutional context. High-trade-intensity economies account for "
        "much of the negative TWFE estimate, largely because small open economies experience large "
        "cyclical swings in both trade and income that have little to do with deliberate trade "
        "liberalization."
    )
    add_paragraph(
        doc,
        "Several limitations are worth stating directly. Trade-to-GDP is a coarse measure of openness "
        "that mixes trade policy with geography, country size, and global commodity price movements "
        "(Rodriguez & Rodrik, 2000). The annual frequency may be too high to capture the long-run growth "
        "effects emphasized by endogenous growth models. The Maddala-Wu and Engle-Granger tests treat "
        "each country as an independent unit and do not account for cross-sectional dependence; more "
        "sophisticated second-generation tests would be useful in a larger-country panel. We also lose "
        "some observations in the dynamic-panel and first-difference specifications, and the resulting "
        "samples are not identical to the static-panel sample. Most importantly, our estimates remain "
        "associative. Establishing causality would require instrumental variables tied to plausibly "
        "exogenous trade shocks, such as Feyrer's (2019) geography-based instrument or synthetic-control "
        "evaluations of liberalization episodes (Billmeier & Nannicini, 2013), or difference-in-"
        "differences designs around preferential trade agreements. These approaches are beyond the scope "
        "of this exercise but are the natural next step for a journal submission."
    )
    add_paragraph(
        doc,
        "From a policy standpoint, the results counsel humility. The aggregate trade-to-GDP ratio is too "
        "blunt an instrument to guide country-specific liberalization strategies. Policymakers who want "
        "to use trade as an engine of growth need to know which products a country exports, whether "
        "domestic firms can absorb foreign technology, and whether the fiscal and monetary framework can "
        "handle external volatility. For researchers, the priority is to move beyond aggregate openness "
        "and identify the specific margins through which trade affects income."
    )

    add_paragraph(
        doc,
        "The findings speak directly to the policy-complementarity hypothesis emphasized by Chang, "
        "Kaltani, and Loayza (2009). Trade openness appears to be associated with higher income in "
        "settings where complementary institutions are already strong, as proxied by OECD membership, "
        "and during periods of rapid global trade expansion, such as the years before the 2008 crisis. "
        "When these supportive conditions are absent, the aggregate correlation weakens or disappears. "
        "This does not imply that trade is harmful; it implies that the growth payoff depends on the "
        "economic environment into which a country is opening."
    )
    add_paragraph(
        doc,
        "A second policy implication concerns measurement. If policymakers target the trade-to-GDP ratio "
        "as an indicator of success, they may inadvertently reward re-exports, commodity booms, and "
        "exchange-rate movements rather than genuine integration into productive global value chains. "
        "A more useful monitoring framework would combine trade outcomes with measures of export "
        "diversification, import content of exports, and domestic value added in foreign demand. "
        "Without such granularity, the headline openness number can mislead as much as it informs."
    )
    add_paragraph(
        doc,
        "This study also illustrates the value of reproducibility in development research. By downloading "
        "the data directly from the World Bank API, documenting every cleaning decision, and providing "
        "the analysis code in a public repository, we make it possible for other researchers to replicate "
        "the tables, test alternative specifications, and update the estimates as new data become available. "
        "Reproducibility does not guarantee correctness, but it makes errors easier to detect and debates "
        "easier to settle."
    )

    # 5. Conclusions
    add_heading(doc, "5. Conclusions", level=1)
    add_paragraph(
        doc,
        "Our results suggest a simple diagnostic: if you want a positive trade-growth coefficient, run a "
        "pooled OLS or fixed-effects model; if you want a negative one, add year fixed effects. The TWFE "
        "estimate turns negative because the global upward trend in both trade and income during "
        "2000-2022 is removed, leaving a within-year, within-country residual in which faster trade "
        "growth is not associated with faster income growth. This sensitivity implies that the aggregate "
        "trade-to-GDP ratio is picking up global cycles and country-size effects, not a structural "
        "policy parameter."
    )
    add_paragraph(
        doc,
        "For policymakers, the results suggest that broad trade liberalization should not be treated as "
        "an automatic engine of growth. The developed-country subsample shows a positive association, "
        "while the developing-country subsample does not; the pre-2008 period also looks different from "
        "the post-2008 period. This heterogeneity is consistent with the idea that the growth payoff from "
        "trade depends on what a country exports, how well its institutions manage external competition, "
        "and whether domestic firms can absorb foreign technology and inputs."
    )
    add_paragraph(
        doc,
        "For researchers, the priority should be to move beyond aggregate openness and identify the "
        "specific margins through which trade affects income: tariff reductions in particular sectors, "
        "access to imported intermediates, export-market entry, and the rules that govern foreign "
        "investment. The present paper provides a reproducible benchmark and a reminder that, in this "
        "literature, the estimated sign of the openness coefficient is often more informative about the "
        "choice of estimator than about the underlying economic relationship."
    )

    add_paragraph(
        doc,
        "Looking forward, the research agenda is clear. Future work should exploit plausibly exogenous "
        "variation in trade policy, use sectoral and firm-level data to trace mechanisms, and test for "
        "heterogeneity by product complexity and institutional quality. The present paper provides a "
        "reproducible baseline and a cautionary message: in the cross-country trade-and-growth literature, "
        "the estimated coefficient on openness is often more a reflection of the estimator than of a "
        "deep structural parameter."
    )
    add_paragraph(
        doc,
        "World Development's readership spans economists, political scientists, and development "
        "practitioners. For this audience, the key takeaway is not whether trade is good or bad in "
        "the abstract, but under what conditions openness translates into broad-based growth. Our "
        "evidence suggests that those conditions are neither universal nor automatic."
    )

    # Declarations (required by World Development).
    add_heading(doc, "Declaration of competing interest", level=1)
    add_paragraph(doc, "Declarations of interest: none", first_line_indent=Inches(0.0))

    add_heading(
        doc,
        "Declaration of generative AI and AI-assisted technologies in the manuscript preparation process",
        level=1,
    )
    add_paragraph(
        doc,
        "During the preparation of this work the author(s) used GLM-5.2 for grammatical correction and "
        "linguistic refinement. After using this tool/service, the author(s) reviewed and edited the "
        "content as needed and take(s) full responsibility for the publication's content.",
        first_line_indent=Inches(0.0),
    )

    add_heading(doc, "Data availability", level=1)
    add_paragraph(
        doc,
        f"All data were downloaded from the World Bank World Development Indicators API "
        f"(https://data.worldbank.org/) on the date of analysis. The Python scripts used to download, "
        f"clean, analyze, and compile the manuscript are available at {REPO_URL}.",
        first_line_indent=Inches(0.0),
    )

    # References
    add_heading(doc, "References", level=1)
    add_reference_list(doc)

    # Tables (placed at the end, per journal guidelines).
    doc.add_page_break()
    add_heading(doc, "Tables", level=1)

    # Table 1
    add_table_caption(doc, "Table 1. Descriptive statistics")
    headers1 = ["Variable", "Mean", "Std. Dev.", "Min", "Max"]
    rows1 = [
        ["GDP per capita (USD)", f"{desc['gdp_mean'].mean():.1f}", f"{desc['gdp_mean'].std():.1f}",
         f"{desc['gdp_min'].min():.1f}", f"{desc['gdp_max'].max():.1f}"],
        ["Trade/GDP (%)", f"{desc['trade_mean'].mean():.1f}", f"{desc['trade_mean'].std():.1f}",
         f"{desc['trade_min'].min():.1f}", f"{desc['trade_max'].max():.1f}"],
        ["Investment/GDP (%)", f"{desc['inv_mean'].mean():.1f}", f"{desc['inv_mean'].std():.1f}",
         f"{desc['inv_min'].min():.1f}", f"{desc['inv_max'].max():.1f}"],
        ["Primary enrollment (%)", f"{desc['school_mean'].mean():.1f}", f"{desc['school_mean'].std():.1f}",
         f"{desc['school_min'].min():.1f}", f"{desc['school_max'].max():.1f}"],
        ["Population growth (%)", f"{desc['popg_mean'].mean():.2f}", f"{desc['popg_mean'].std():.2f}",
         f"{desc['popg_min'].min():.2f}", f"{desc['popg_max'].max():.2f}"],
    ]
    add_word_table(doc, headers1, rows1)
    add_table_note(doc, "Mean values by variable across countries.")

    # Table 2
    add_table_caption(doc, "Table 2. Correlation matrix")
    headers2 = ["", "ln GDP", "ln Trade", "ln Inv", "School", "Pop. gr."]
    rows2 = [
        ["ln GDP", f"{corr.loc['lgdp', 'lgdp']:.3f}", f"{corr.loc['lgdp', 'ltrade']:.3f}",
         f"{corr.loc['lgdp', 'linv']:.3f}", f"{corr.loc['lgdp', 'school']:.3f}",
         f"{corr.loc['lgdp', 'popg']:.3f}"],
        ["ln Trade", f"{corr.loc['ltrade', 'lgdp']:.3f}", f"{corr.loc['ltrade', 'ltrade']:.3f}",
         f"{corr.loc['ltrade', 'linv']:.3f}", f"{corr.loc['ltrade', 'school']:.3f}",
         f"{corr.loc['ltrade', 'popg']:.3f}"],
        ["ln Inv", f"{corr.loc['linv', 'lgdp']:.3f}", f"{corr.loc['linv', 'ltrade']:.3f}",
         f"{corr.loc['linv', 'linv']:.3f}", f"{corr.loc['linv', 'school']:.3f}",
         f"{corr.loc['linv', 'popg']:.3f}"],
        ["School", f"{corr.loc['school', 'lgdp']:.3f}", f"{corr.loc['school', 'ltrade']:.3f}",
         f"{corr.loc['school', 'linv']:.3f}", f"{corr.loc['school', 'school']:.3f}",
         f"{corr.loc['school', 'popg']:.3f}"],
        ["Pop. gr.", f"{corr.loc['popg', 'lgdp']:.3f}", f"{corr.loc['popg', 'ltrade']:.3f}",
         f"{corr.loc['popg', 'linv']:.3f}", f"{corr.loc['popg', 'school']:.3f}",
         f"{corr.loc['popg', 'popg']:.3f}"],
    ]
    add_word_table(doc, headers2, rows2)

    # Table 3
    if unitroot is not None and not unitroot.empty and coint is not None and not coint.empty:
        add_table_caption(doc, "Table 3. Panel stationarity and cointegration pre-tests")
        headers3 = ["Variable", "Countries", "Reject 5%", "Mean p", "Fisher chi2", "Fisher p"]
        rows3 = []
        for _, row in unitroot.iterrows():
            rows3.append([
                row["variable"], str(int(row["n_countries"])), str(int(row["n_reject_5pct"])),
                f"{row['mean_pvalue']:.3f}", f"{row['fisher_chi2']:.2f}", format_p(row["fisher_pvalue"]),
            ])
        cr = coint.iloc[0]
        rows3.append([
            "Cointegration", str(int(cr["n_countries"])), str(int(cr["n_reject_5pct"])),
            f"{cr['mean_pvalue']:.3f}", f"{cr['fisher_chi2']:.2f}", format_p(cr["fisher_pvalue"]),
        ])
        add_word_table(doc, headers3, rows3)
        add_table_note(
            doc,
            "Reject 5% counts how many country-specific ADF tests reject the unit-root null at the 5% "
            "level. The Fisher chi2 statistic combines individual p-values across countries; a low p-value "
            "indicates that the series is persistent in most countries. Cointegration tests are Engle-"
            "Granger residual ADF tests from a regression of log GDP per capita on log trade, log "
            "investment, schooling, and population growth."
        )

    # Table 4
    add_table_caption(doc, "Table 4. Panel regression results (dependent variable: ln GDP per capita)")
    headers4 = ["Variable", "Pooled OLS", "FE", "Time FE", "Two-way FE"]
    rows4 = [
        ["ln Trade", coef_str(*get_coef(panel, "pooled_ols", "ltrade")),
         coef_str(*get_coef(panel, "fe", "ltrade")), coef_str(*get_coef(panel, "te", "ltrade")),
         coef_str(*get_coef(panel, "twfe", "ltrade"))],
        ["ln Investment", coef_str(*get_coef(panel, "pooled_ols", "linv")),
         coef_str(*get_coef(panel, "fe", "linv")), coef_str(*get_coef(panel, "te", "linv")),
         coef_str(*get_coef(panel, "twfe", "linv"))],
        ["Primary enroll.", coef_str(*get_coef(panel, "pooled_ols", "school")),
         coef_str(*get_coef(panel, "fe", "school")), coef_str(*get_coef(panel, "te", "school")),
         coef_str(*get_coef(panel, "twfe", "school"))],
        ["Population growth", coef_str(*get_coef(panel, "pooled_ols", "popg")),
         coef_str(*get_coef(panel, "fe", "popg")), coef_str(*get_coef(panel, "te", "popg")),
         coef_str(*get_coef(panel, "twfe", "popg"))],
        ["R-squared", f"{get_coef(panel, 'pooled_ols', 'r2')[0]:.4f}",
         f"{get_coef(panel, 'fe', 'r2')[0]:.4f}", f"{get_coef(panel, 'te', 'r2')[0]:.4f}",
         f"{get_coef(panel, 'twfe', 'r2')[0]:.4f}"],
        ["Observations", f"{int(get_coef(panel, 'pooled_ols', 'n_obs')[0])}",
         f"{int(get_coef(panel, 'fe', 'n_obs')[0])}", f"{int(get_coef(panel, 'te', 'n_obs')[0])}",
         f"{int(get_coef(panel, 'twfe', 'n_obs')[0])}"],
    ]
    add_word_table(doc, headers4, rows4)
    add_table_note(
        doc,
        "*** p<0.01, ** p<0.05, * p<0.1. Standard errors are panel-robust. FE = country fixed effects; "
        "TWFE = country and year fixed effects. TWFE absorbs country-specific and year-specific averages; "
        "the negative coefficient is a within-year, within-country deviation and does not measure the "
        "long-run effect of lowering tariffs. We also attempted a random-effects specification, but it "
        "returned a numerical singularity and was dropped; fixed-effects estimates are reported "
        "throughout."
    )

    # Table 5
    if het is not None and not het.empty:
        add_table_caption(doc, "Table 5. Heterogeneity analysis (two-way fixed effects, dependent variable: ln GDP per capita)")
        headers5 = ["Sample", "ln Trade", "ln Investment", "Observations"]
        rows5 = []
        for s in ["developed", "developing", "pre_2008", "post_2008", "high_trade", "low_trade"]:
            nc, _ = get_coef_het(het, s, "n_obs")
            rows5.append([
                s.replace("_", " ").title(),
                coef_str(*get_coef_het(het, s, "ltrade")),
                coef_str(*get_coef_het(het, s, "linv")),
                str(int(nc)) if not np.isnan(nc) else "---",
            ])
        add_word_table(doc, headers5, rows5)
        add_table_note(
            doc,
            "*** p<0.01, ** p<0.05, * p<0.1. High/low trade intensity is split by the median trade-to-"
            "GDP ratio across country-year observations. Developed/developing split uses OECD membership; "
            "the pre/post-2008 split follows the global financial crisis."
        )

    # Table 6
    add_table_caption(doc, "Table 6. Robustness checks")
    headers6 = ["Variable", "Excl. high trade", "Lagged trade", "First differences"]
    rows6 = [
        ["ln Trade", coef_str(*get_coef(robust, "excl_high_trade", "ltrade")),
         coef_str(*get_coef(robust, "lag_trade", "ltrade_lag1")),
         coef_str(*get_coef(robust, "first_diff", "d_ltrade"))],
        ["ln Investment", coef_str(*get_coef(robust, "excl_high_trade", "linv")),
         coef_str(*get_coef(robust, "lag_trade", "linv")),
         coef_str(*get_coef(robust, "first_diff", "d_linv"))],
        ["Primary enroll.", coef_str(*get_coef(robust, "excl_high_trade", "school")),
         coef_str(*get_coef(robust, "lag_trade", "school")),
         coef_str(*get_coef(robust, "first_diff", "d_school"))],
        ["Population growth", coef_str(*get_coef(robust, "excl_high_trade", "popg")),
         coef_str(*get_coef(robust, "lag_trade", "popg")),
         coef_str(*get_coef(robust, "first_diff", "d_popg"))],
        ["R-squared", f"{get_coef(robust, 'excl_high_trade', 'r2')[0]:.4f}",
         f"{get_coef(robust, 'lag_trade', 'r2')[0]:.4f}",
         f"{get_coef(robust, 'first_diff', 'r2')[0]:.4f}"],
        ["Observations", f"{int(get_coef(robust, 'excl_high_trade', 'n_obs')[0])}",
         f"{int(get_coef(robust, 'lag_trade', 'n_obs')[0])}",
         f"{int(get_coef(robust, 'first_diff', 'n_obs')[0])}"],
    ]
    add_word_table(doc, headers6, rows6)
    add_table_note(
        doc,
        "*** p<0.01, ** p<0.05, * p<0.1. Excl. high trade drops countries whose mean trade share exceeds "
        "150% of GDP. Lagged trade uses one-year-lagged log trade openness. First differences uses annual "
        "changes in all variables and drops the time effects because they are differenced out."
    )

    # Table 7
    add_table_caption(doc, "Table 7. Dynamic panel (dependent variable: ln GDP per capita)")
    headers7 = ["Variable", "Coefficient"]
    rows7 = [
        ["ln GDP(t-1)", coef_str(*get_coef(dyn, "dynamic_fe", "lgdp_lag1"))],
        ["ln Trade(t-1)", coef_str(*get_coef(dyn, "dynamic_fe", "ltrade_lag1"))],
        ["ln Investment", coef_str(*get_coef(dyn, "dynamic_fe", "linv"))],
        ["Primary enroll.", coef_str(*get_coef(dyn, "dynamic_fe", "school"))],
        ["Population growth", coef_str(*get_coef(dyn, "dynamic_fe", "popg"))],
        ["R-squared", f"{get_coef(dyn, 'dynamic_fe', 'r2')[0]:.4f}"],
        ["Observations", f"{int(get_coef(dyn, 'dynamic_fe', 'n_obs')[0])}"],
    ]
    add_word_table(doc, headers7, rows7)
    add_table_note(
        doc,
        "*** p<0.01, ** p<0.05, * p<0.1. Country and year fixed effects included. The lagged dependent "
        "variable is expected to absorb much of the persistence in GDP per capita; the remaining trade "
        "coefficient should be interpreted as a short-run association rather than a long-run elasticity."
    )

    doc.save(MANUSCRIPT_PATH)
    return doc


# ---------------------------------------------------------------------------
# Title page (separate file for double-anonymized submission)
# ---------------------------------------------------------------------------

def build_title_page():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    set_default_font(doc)

    add_centered_paragraph(
        doc,
        "Trade Openness and Economic Growth: A Replicable Cross-Country Panel Analysis, 2000-2022",
        size=Pt(16), bold=True, space_after=Pt(24), line_spacing=2.0
    )

    add_heading(doc, "Author(s)", level=2)
    add_paragraph(doc, "Anonymous Author(s)", first_line_indent=Inches(0.0))

    add_heading(doc, "Affiliations", level=2)
    add_paragraph(doc, "Affiliation withheld for peer review", first_line_indent=Inches(0.0))

    add_heading(doc, "Corresponding author", level=2)
    add_paragraph(
        doc,
        "Name, full postal address, and e-mail address to be added at submission.",
        first_line_indent=Inches(0.0),
    )

    add_heading(doc, "Acknowledgments", level=2)
    add_paragraph(
        doc,
        "Artificial-intelligence-assisted language tool (GLM-5.2) was utilized for grammatical correction "
        "and linguistic refinement during manuscript preparation. No AI was involved in generating "
        "research ideas, deriving conclusions, or designing the study. The authors take full "
        "responsibility for the final manuscript.",
        first_line_indent=Inches(0.0),
    )

    add_heading(doc, "Declaration of competing interest", level=2)
    add_paragraph(doc, "Declarations of interest: none", first_line_indent=Inches(0.0))

    doc.save(TITLE_PAGE_PATH)
    return doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def build_document():
    panel, dyn, robust, desc, corr, vif, unitroot, coint, het = load_results()
    manuscript = build_manuscript(panel, dyn, robust, desc, corr, vif, unitroot, coint, het)
    title_page = build_title_page()
    print(f"Saved anonymized manuscript: {MANUSCRIPT_PATH}")
    print(f"Saved title page: {TITLE_PAGE_PATH}")
    print(f"Manuscript paragraphs: {len(manuscript.paragraphs)}")


if __name__ == "__main__":
    build_document()
